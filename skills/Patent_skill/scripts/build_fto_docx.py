#!/usr/bin/env python3
"""Build a template-derived, evidence-bounded FTO report DOCX."""

import argparse
import csv
import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Mm, Pt, RGBColor


# Arial Unicode MS is present in the desktop LibreOffice environment and
# renders both simplified Chinese and Latin patent identifiers consistently.
FONT = "STSong"
FONT_FALLBACK = "Arial Unicode MS"
NAVY = "17365D"
BLUE = "2E74B5"
LIGHT_BLUE = "D9EAF7"
PALE_BLUE = "EEF5FB"
MUTED = "667085"
INK = "1F2937"
CAUTION = "8A5A00"
RED = "9B1C1C"
GREEN = "287A50"
PORTRAIT_WIDTH = 8300
LANDSCAPE_WIDTH = 13950


def load_json(path):
    return json.loads(path.read_text(encoding="utf-8"))


def csv_rows(path):
    with path.open(newline="", encoding="utf-8-sig") as handle:
        return [{str(k).strip(): (v or "").strip() for k, v in row.items()} for row in csv.DictReader(handle)]


def set_run_font(run, size=10.5, color=INK, bold=False, italic=False, name=FONT):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), name)
    rfonts.set(qn("w:hint"), "eastAsia")


def style_paragraph(style, size, color=INK, bold=False, before=0, after=6, line=1.15):
    style.font.name = FONT
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = line
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), FONT)
    rfonts.set(qn("w:hint"), "eastAsia")


def ensure_style(doc, name, size, color=INK, bold=False, before=0, after=6, line=1.15):
    try:
        style = doc.styles[name]
    except KeyError:
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style_paragraph(style, size, color, bold, before, after, line)
    return style


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=0):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[min(idx, len(widths) - 1)]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_borders(table, color="B8C9D9", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def clear_cell(cell):
    cell.text = ""
    return cell.paragraphs[0]


def add_cell_text(cell, text, size=9, color=INK, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    p = clear_cell(cell)
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(str(text or "—"))
    set_run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_hyperlink(paragraph, text, url, size=8.5, color=BLUE):
    part = paragraph.part
    rid = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), FONT)
    rfonts.set(qn("w:hint"), "eastAsia")
    rpr.append(rfonts)
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), color)
    rpr.append(color_node)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rpr.append(sz)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_field(paragraph, field):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(text)
    run._r.append(end)
    set_run_font(run, size=8, color=MUTED)


def set_page(section, landscape=False):
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Mm(297)
        section.page_height = Mm(210)
        section.left_margin = Mm(25.4)
        section.right_margin = Mm(25.4)
        section.top_margin = Mm(31.75)
        section.bottom_margin = Mm(31.75)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.left_margin = Mm(31.75)
        section.right_margin = Mm(31.75)
        section.top_margin = Mm(25.4)
        section.bottom_margin = Mm(25.4)
    section.header_distance = Mm(12.5)
    section.footer_distance = Mm(12.5)


def set_footer(section, date_text):
    section.footer.is_linked_to_previous = False
    footer = section.footer
    p = footer.paragraphs[0]
    for run in list(p.runs):
        p._p.remove(run._r)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"FTO · {date_text} · ")
    set_run_font(r, size=8, color=MUTED)
    add_field(p, "PAGE")


def add_para(doc, text="", size=10.5, color=INK, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=6, line=1.15):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_labeled_para(doc, label, text, after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(label + "：")
    set_run_font(r, size=10.5, color=NAVY, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)
    return p


def add_heading(doc, text, level=1):
    style_name = "一级标题" if level == 1 else "二级标题"
    p = doc.add_paragraph(style=style_name)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(r, size=16 if level == 1 else 12.5, color=BLUE if level == 1 else NAVY, bold=True)
    return p


def add_callout(doc, label, text, fill=PALE_BLUE, color=NAVY):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [PORTRAIT_WIDTH])
    set_table_borders(table, color="BFD5E8", size="8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = clear_cell(cell)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(label + "  ")
    set_run_font(r, size=10, color=color, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths, header_fill=LIGHT_BLUE, body_size=8.8):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    repeat_header(table.rows[0])
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, header_fill)
        add_cell_text(cell, header, size=9, color=NAVY, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            add_cell_text(cells[i], value, size=body_size, color=INK)
    return table


def compact(value, limit=380):
    value = " ".join(str(value or "").split())
    return value if len(value) <= limit else value[: limit - 1] + "…"


def case_descriptor(scope, identity):
    obj = scope.get("research_object", {})
    molecule = obj.get("molecule") or identity.get("molecule", {}).get("canonical") or "研究对象"
    target = obj.get("target") or identity.get("target", {}).get("canonical") or "目标机制"
    indication = obj.get("indication") or identity.get("indication", {}).get("canonical") or "目标用途"
    return molecule, target, indication


def case_feature_bullets(plan, limit=4):
    items = [compact(feature.get("text"), 260) for feature in plan.get("features", []) if feature.get("text")]
    return items[:limit] or ["当前案例未提供可结构化的技术特征；请先补充 fto-input.json。"]


def build_report(project):
    plan = load_json(project / "fto-search-plan.json")
    candidates = csv_rows(project / "fto-candidate-ranking.csv")
    families = csv_rows(next(project.glob("*-patent-families.csv")))
    claims = csv_rows(next(project.glob("*-claim-elements.csv")))
    evidence = csv_rows(next(project.glob("*-evidence.csv")))
    family_map = {row.get("family_id"): row for row in families}
    candidate_map = {row.get("family_id"): row for row in candidates}
    claim_groups = {}
    for row in claims:
        claim_groups.setdefault(row.get("family_id"), []).append(row)

    scope = plan.get("scope", {})
    source_catalog = plan.get("source_catalog", {})
    source_counts = source_catalog.get("counts", {})
    as_of = scope.get("as_of", "未注明")
    molecule, target, indication = case_descriptor(scope, plan.get("entity_resolution", {}))
    count_by_priority = {key: sum(1 for row in candidates if row.get("review_priority") == key) for key in ("HIGH", "MEDIUM", "LOW")}
    doc = Document()
    # Source-derived A4 system: portrait narrative, landscape comparison, portrait appendices.
    for section in doc.sections:
        set_page(section, landscape=False)
        set_footer(section, as_of)
    ensure_style(doc, "Normal", 10.5, INK, False, 0, 6, 1.15)
    ensure_style(doc, "一级标题", 16, BLUE, True, 14, 8, 1.1)
    ensure_style(doc, "二级标题", 12.5, NAVY, True, 9, 5, 1.1)
    ensure_style(doc, "封面标题", 24, NAVY, True, 0, 8, 1.1)
    # Built-in list styles are configured so wrapped lines retain a real list indent.
    for name in ("List Bullet", "List Number"):
        try:
            style_paragraph(doc.styles[name], 10.5, INK, False, 0, 4, 1.15)
            doc.styles[name].paragraph_format.left_indent = Mm(9)
            doc.styles[name].paragraph_format.first_line_indent = Mm(-4.5)
        except KeyError:
            pass

    # Cover.
    add_para(doc, "FTO 初筛报告", size=11, color=BLUE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=54, after=14)
    p = doc.add_paragraph(style="封面标题")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(compact(plan.get("technical_solution") or f"{molecule} 相关技术方案", 72))
    set_run_font(r, size=23, color=NAVY, bold=True)
    add_para(doc, "FTO（Freedom to Operate）公开专利初筛与权利要求要素比对报告", size=12, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=28)
    add_para(doc, f"报告基准日：{as_of}  ·  研究区域：{', '.join(scope.get('jurisdictions', []) + scope.get('related_jurisdictions', []))}", size=10.5, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER, after=56)
    metric_rows = [["检索策略轮次", "候选专利族", "中优先复核", "权利要求要素记录"], ["7（计划）", str(len(candidates)), str(count_by_priority["MEDIUM"]), str(len(claims))]]
    table = add_table(doc, metric_rows[0], [metric_rows[1]], [2075, 2075, 2075, 2075], header_fill=LIGHT_BLUE, body_size=12)
    for cell in table.rows[1].cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            set_run_font(run, size=13, color=BLUE, bold=True)
    add_para(doc, "本报告基于当前案例目录中的公开专利镜像、专利族记录和权利要求要素记录生成；未将排序分数表述为侵权概率或法律结论。", size=9, color=MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=26, after=0)
    doc.add_page_break()

    add_heading(doc, "概览")
    top_candidates = [row for row in candidates if row.get("review_priority") in ("HIGH", "MEDIUM")]
    top_label = "、".join(f"{row.get('family_id')} / {row.get('representative_document')}" for row in top_candidates[:3]) or "当前候选族"
    add_callout(doc, "初筛结论", f"在现有公开专利数据和已记录权利要求要素范围内，尚不足以确认拟实施方案落入某一目标法域的有效独立权利要求。优先复核对象为 {top_label}；应回到完整独立权利要求、国家阶段和官方状态逐项核验。")
    add_para(doc, "本报告的“高/中/低”是候选专利的复核优先级，不是侵权风险等级。核心组合物族、制剂族和邻近生物标志物族仍需分别核对其目标法域的成员、独立权利要求、分案/继续申请和官方法律状态。", after=8)
    for text in [
        f"技术方案围绕 {molecule}、{target} 和 {indication} 展开。",
        f"本轮已结构化 {len(plan.get('search_rounds', []))} 个检索轮次、{len(plan.get('keyword_expansion', []))} 个关键词簇、{len(plan.get('classifications', []))} 个 IPC/CPC 候选号，并对 {len(candidates)} 个候选专利族、{len(claims)} 条权利要求要素记录进行透明排序。",
        f"来源目录已纳入 CNIPA/PatentDatabases 的 {source_counts.get('upstream_listings', '—')} 条记录（去重后 {source_counts.get('unique_urls', '—')} 个 URL）；本案例候选证据主要来自 Google Patents 公共镜像，官方状态核验仍需按目标法域逐项完成。",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(text)
        set_run_font(r, size=10.5, color=INK)

    add_heading(doc, "检索范围")
    add_labeled_para(doc, "研究对象", f"{molecule}；靶点/机制：{target}；适应症/用途：{indication}。")
    add_labeled_para(doc, "目标法域", f"{', '.join(scope.get('jurisdictions', []))}；关联扩展法域 {', '.join(scope.get('related_jurisdictions', []))}。")
    add_labeled_para(doc, "时间边界", f"截至 {as_of} 的公开记录；未公开申请、未核验的国家阶段或状态变化不在本报告中假定。")
    add_labeled_para(doc, "数据范围", "专利族 CSV、权利要求要素 CSV、证据链 CSV、来源日志及 FTO 检索计划。")
    add_labeled_para(doc, "来源目录", f"CNIPA/PatentDatabases 快照：上游 {source_counts.get('upstream_listings', '—')} 条记录、去重后 {source_counts.get('unique_urls', '—')} 个 URL；目录用于选择检索入口，不代表所有来源已在本案例中实际访问。")
    add_labeled_para(doc, "检索限制", "当前案例是标准分析级别的证据包，不等同于穷尽式检索或正式律师 FTO 意见。")

    add_heading(doc, "技术方案")
    add_para(doc, plan.get("technical_solution", ""), after=8)
    add_para(doc, "拟实施方案中的结构化技术特征包括：", color=NAVY, bold=True, after=3)
    for text in case_feature_bullets(plan):
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(text)
        set_run_font(r, size=10.5, color=INK)

    add_heading(doc, "技术特征列表")
    feature_rows = []
    for feature in plan.get("features", []):
        feature_rows.append([feature.get("id"), feature.get("feature_type"), feature.get("importance"), compact(feature.get("text"), 260), ", ".join(feature.get("classifications", [])) or "—"])
    add_table(doc, ["编号", "类型", "重要性", "技术特征", "IPC/CPC"], feature_rows, [700, 1200, 1050, 3900, 1450], body_size=8.8)

    add_heading(doc, "专利族初筛总览")
    overview_rows = []
    for row in candidates:
        family = family_map.get(row.get("family_id"), {})
        overview_rows.append([
            row.get("family_id"),
            row.get("representative_document"),
            family.get("claim_theme", row.get("relevance")),
            row.get("review_priority"),
            f"{float(row.get('screen_score') or 0):.2f}",
            row.get("matched_features") or "—",
            compact(family.get("official_status", ""), 150),
        ])
    add_table(doc, ["专利族", "代表文献", "主题", "复核优先级", "初筛分数", "完整命中特征", "状态信号"], overview_rows, [900, 1300, 2750, 1050, 800, 1100, 1400], body_size=8.2)
    add_para(doc, "注：分数由声明的关键词簇覆盖、权利要求类别、族相关性和状态信号组成；它用于排序，不代表权利要求覆盖概率。", size=8.5, color=MUTED, italic=True, before=4, after=10)

    # Landscape claim comparison section.
    landscape = doc.add_section(WD_SECTION.NEW_PAGE)
    set_page(landscape, landscape=True)
    set_footer(landscape, as_of)
    add_heading(doc, "技术特征比对表")
    add_para(doc, "本表把已有 claim-elements 记录与 FTO 技术特征进行对照。由于当前记录不是完整的逐项独立权利要求文本，表内使用“明确披露/可能覆盖/部分命中/待核验”等证据语言，不使用“必然侵权”等表述。", after=8)
    comparison_rows = []
    feature_text = {f.get("id"): f.get("text") for f in plan.get("features", [])}
    for claim in claims:
        ranking = candidate_map.get(claim.get("family_id"), {})
        family = family_map.get(claim.get("family_id"), {})
        coverage = claim.get("coverage", "")
        if coverage == "明确披露":
            signal = "重点复核"
            signal_color = RED
        elif coverage == "可能覆盖":
            signal = "部分命中"
            signal_color = CAUTION
        else:
            signal = "待核验"
            signal_color = MUTED
        comparison_rows.append([
            f"{claim.get('family_id')}\n{claim.get('document')}",
            compact(claim.get("element"), 330),
            compact(ranking.get("matched_features") or ranking.get("partial_features") or "—", 100),
            claim.get("claim_category") or "—",
            signal,
            compact(claim.get("notes") or family.get("notes") or "需核对完整独立权利要求和官方状态。", 270),
        ])
    table = add_table(doc, ["专利族/文献", "权利要求要素记录", "对应特征", "权利要求类别", "初筛信号", "判定理由/下一步"], comparison_rows, [1450, 3900, 1150, 1200, 1050, 5200], body_size=8.1)
    # Color signal cells without relying on unverified legal language.
    for idx, claim in enumerate(claims, start=1):
        cell = table.rows[idx].cells[4]
        value = comparison_rows[idx - 1][4]
        color = RED if value == "重点复核" else CAUTION if value == "部分命中" else MUTED
        for run in cell.paragraphs[0].runs:
            set_run_font(run, size=8.1, color=color, bold=True)

    add_heading(doc, "FTO 复核优先级矩阵")
    matrix_rows = []
    for row in candidates:
        family = family_map.get(row.get("family_id"), {})
        matrix_rows.append([
            row.get("family_id"),
            row.get("review_priority"),
            family.get("relevance", row.get("relevance")),
            row.get("matched_features") or "—",
            row.get("partial_features") or "—",
            compact(family.get("status_source", ""), 160),
        ])
    add_table(doc, ["专利族", "优先级", "主集合/边界", "完整命中", "部分命中", "状态证据"], matrix_rows, [1400, 1000, 1400, 1700, 2200, 6250], body_size=8.5)
    add_para(doc, "当前最值得先做 claim chart 和官方状态核验的是 DVL-FAM-004；但“中优先复核”不等于“已确认侵权”。", size=9, color=CAUTION, italic=True, before=5, after=8)

    # Portrait details and appendices.
    portrait = doc.add_section(WD_SECTION.NEW_PAGE)
    set_page(portrait, landscape=False)
    set_footer(portrait, as_of)
    add_heading(doc, "重点专利族与证据链")
    for family_id, family in family_map.items():
        ranking = candidate_map.get(family_id, {})
        add_heading(doc, f"{family_id} · {family.get('representative_document', '—')}", level=2)
        detail_rows = [
            ["技术主题", compact(family.get("claim_theme"), 260)],
            ["代表文献", family.get("representative_document")],
            ["优先级", f"{ranking.get('review_priority', '—')} · 初筛分数 {ranking.get('screen_score', '—')}"],
            ["权利要求类别", family.get("claim_categories")],
            ["关键要素", compact(family.get("key_claim_elements"), 300)],
            ["完整/部分命中", f"{ranking.get('matched_features') or '无'} / {ranking.get('partial_features') or '无'}"],
            ["状态信号", compact(family.get("official_status"), 250)],
        ]
        add_table(doc, ["字段", "内容"], detail_rows, [1800, 6500], body_size=8.8, header_fill="F2F4F7")
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run("FTO 初筛解释：")
        set_run_font(r, size=9.5, color=NAVY, bold=True)
        if ranking.get("review_priority") in ("HIGH", "MEDIUM"):
            text = "该族出现了案例声明的核心或必要技术特征信号，建议先核验目标法域国家阶段、独立权利要求和审查历史；现有初筛不能确认所有必要特征被同一独立权利要求组合限定。"
        elif family.get("relevance") == "boundary":
            text = "该族属于邻近技术或竞争边界，当前未建立对拟实施方案的完整权利要求联系；保留作为边界线索，需以独立权利要求和法域成员复核。"
        else:
            text = "当前特征工程只形成部分或背景层命中，尚不足以作为实施障碍结论；仍应检查同族分支、国家阶段、分案/继续申请及权利要求变化。"
        r = p.add_run(text)
        set_run_font(r, size=9.5, color=INK)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        r = p.add_run("来源：")
        set_run_font(r, size=8.8, color=MUTED, bold=True)
        add_hyperlink(p, family.get("source_url", ""), family.get("source_url", ""), size=8.2)

    add_heading(doc, "证据链摘要")
    evidence_rows = []
    for item in evidence:
        evidence_rows.append([item.get("finding_id"), compact(item.get("conclusion_or_fact"), 240), item.get("evidence_type"), compact(item.get("claim_or_event_location"), 130), item.get("confidence"), item.get("source_url")])
    add_table(doc, ["Finding", "事实/结论", "证据类型", "定位", "置信度", "来源"], evidence_rows, [800, 3300, 1200, 1500, 850, 650], body_size=8.2)

    add_heading(doc, "附录 A：关键词扩展与 IPC/CPC")
    keyword_rows = []
    for cluster in plan.get("keyword_expansion", []):
        linked = ", ".join(feature.get("id") for feature in plan.get("features", []) if cluster.get("id") in feature.get("keyword_clusters", [])) or "—"
        keyword_rows.append([cluster.get("label"), ", ".join(cluster.get("base_terms", [])), compact(", ".join(cluster.get("expanded_terms", [])), 300), linked, cluster.get("source", "")])
    add_table(doc, ["关键词簇", "基础词", "扩展词", "关联特征", "来源"], keyword_rows, [1200, 1900, 3650, 900, 650], body_size=8.1)
    add_heading(doc, "IPC/CPC 候选号")
    add_para(doc, ", ".join(plan.get("classifications", [])), size=9, color=INK, after=8)
    add_para(doc, "分类号是提高召回和定位邻近技术的入口，不代表分类号覆盖了拟实施方案；正式 FTO 应以命中文献的权利要求和目标法域官方记录为准。", size=8.8, color=MUTED, italic=True)

    add_heading(doc, "附录 B：防侵权检索策略")
    search_rows = []
    for row in plan.get("search_rounds", []):
        search_rows.append([row.get("id"), row.get("title"), ", ".join(row.get("fields", [])), compact(row.get("objective"), 210), compact(row.get("formula"), 500)])
    add_table(doc, ["轮次", "类型", "检索字段", "目标", "检索式/关系扩展"], search_rows, [650, 1700, 1350, 2700, 1900], body_size=7.8)
    add_para(doc, "说明：本案例中的检索轮次是可恢复的策略计划；每轮真实命中数量、数据库、日期、纳排决定和来源 URL应继续写入 source-log.jsonl。", size=8.8, color=MUTED, italic=True, before=4, after=8)

    add_heading(doc, "附录 C：来源目录与使用原则")
    add_para(doc, "本 Skill 将 CNIPA/PatentDatabases 作为来源发现目录，并按来源角色分层使用。完整名称、URL、上游分组、来源角色和快照哈希保存在案例目录的 patent-database-sources.json；本表只列出路由统计。", size=8.8, color=INK, after=6)
    source_summary = []
    for source_kind, count in sorted((source_counts.get("by_source_kind") or {}).items()):
        source_summary.append([source_kind, str(count), "官方/登记簿核验" if source_kind == "official_or_authority" else "来源发现与交叉核对" if source_kind in ("commercial_or_aggregator", "public_or_national_database") else "技术/临床上下文" if source_kind == "literature_or_context" else "分类号导航与版本核对"])
    add_table(doc, ["来源角色", "数量", "本 Skill 的使用原则"], source_summary, [2500, 900, 4900], body_size=8.6)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("上游目录：")
    set_run_font(r, size=8.8, color=MUTED, bold=True)
    add_hyperlink(p, source_catalog.get("upstream_repo", "https://github.com/CNIPA/PatentDatabases"), source_catalog.get("upstream_repo", "https://github.com/CNIPA/PatentDatabases"), size=8.2)

    add_heading(doc, "免责声明")
    add_para(doc, "本报告为基于当前案例公开专利资料的 FTO 初步检索与权利要求要素比对材料，不构成侵权、有效性、不侵权或自由实施的法律意见。报告中的“重点复核、中优先、低优先、部分命中”等仅表示检索和人工复核顺序，不代表侵权概率。涉及商业实施、上市、许可、诉讼或重大投资决策时，应由具备相关法域资格的专利律师核对完整独立权利要求、审查档案、同族/分案/继续申请和官方法律状态。聚合网站的状态标签不能替代 CNIPA、USPTO、WIPO、EPO 或其他目标法域官方登记簿。", size=9.2, after=8)
    add_para(doc, f"生成日期：{datetime.now().strftime('%Y-%m-%d')} · 案例目录：{project.name}", size=8.5, color=MUTED, italic=True)
    return doc


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--project-dir", required=True)
    parser.add_argument("--output", default="fto-screening-report.docx")
    args = parser.parse_args()
    project = Path(args.project_dir).expanduser().resolve()
    output = project / args.output
    doc = build_report(project)
    doc.core_properties.title = f"{project.name} FTO 初筛报告"
    doc.core_properties.subject = "公开专利初筛与权利要求要素比对"
    doc.core_properties.author = "Medtech Patent Roadmap"
    doc.save(output)
    print(f"Generated {output}")


if __name__ == "__main__":
    main()
